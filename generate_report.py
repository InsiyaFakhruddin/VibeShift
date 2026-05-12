from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ─────────────────────────────────────────────────────────────
section = doc.sections[0]
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)

# ── Style helpers ─────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(16)
    return p

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(14)
    return p

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(12)
    return p

def para(text, bold_prefix=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
    p.add_run(text)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.bold = True
    p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_after = Pt(4)
    return p

# ═══════════════════════════════════════════════════════════════════════════════
# CHAPTER 5
# ═══════════════════════════════════════════════════════════════════════════════
h1("Chapter 5: Implementation")

# ── 5.1 ──────────────────────────────────────────────────────────────────────
h2("5.1  Chapter Overview")
para(
    "This chapter presents the technical implementation of VibeShift, a mobile application "
    "for AI-driven music transformation. It details the technology stack selection rationale, "
    "frontend and backend architecture, AI microservice integration via the Replicate API, "
    "asynchronous job processing, database schema design, and cloud storage configuration. "
    "The system follows a client-server architecture in which a React Native mobile client "
    "communicates with a Python FastAPI backend, which in turn delegates computationally "
    "intensive AI workloads—stem separation and genre transformation—to Replicate's cloud-hosted "
    "inference platform, eliminating the need for self-managed GPU infrastructure."
)

# ── 5.2 ──────────────────────────────────────────────────────────────────────
h2("5.2  Technology Stack")

h3("5.2.1  Frontend — React Native / Expo")
para(
    "The mobile client is built with React Native 0.81.5 and the Expo framework (v54.0.27), "
    "enabling cross-platform deployment to iOS and Android from a single TypeScript 5.9.2 "
    "codebase. Expo Router 6.0.17 provides file-based navigation. Key dependencies include:"
)
bullet("@clerk/clerk-expo ^2.19.31", bold_prefix="Authentication")
bullet("react-native-svg 15.12.1", bold_prefix="Stem ring visualisation")
bullet("react-native-reanimated ~4.1.1", bold_prefix="Gesture & animation")
bullet("expo-linear-gradient", bold_prefix="UI gradients")
bullet("expo-secure-store", bold_prefix="Secure JWT token storage")
bullet("expo-av", bold_prefix="Audio playback")

h3("5.2.2  Backend API — Python / FastAPI")
para(
    "The REST API is implemented in FastAPI 0.104.1, served by Uvicorn 0.24.0 under Gunicorn "
    "as the process manager (300-second worker timeout to accommodate long-running AI calls). "
    "SQLModel 0.0.19—a hybrid of Pydantic and SQLAlchemy—serves as the ORM. Additional key "
    "libraries:"
)
bullet("httpx — async HTTP client for Replicate API calls")
bullet("boto3 — AWS S3 integration")
bullet("PyJWT[crypto] 2.8.0 — RS256 JWT verification")
bullet("python-multipart — multipart/form-data file upload handling")

h3("5.2.3  AI Microservices — Python / Replicate API")
para(
    "Rather than maintaining self-hosted GPU infrastructure, VibeShift delegates all AI inference "
    "to Replicate, a managed cloud platform for running machine learning models on-demand. Two "
    "models are integrated:"
)
bullet(
    "cjwbw/demucs:25a173108cff36ef9f80f854c162d01df9e6528be175794b81158fa03836d953",
    bold_prefix="Demucs (stem separation)"
)
bullet(
    "meta/musicgen:7a76a8258b23fae65c5a22debb8841d1d7e816b75c2f24218cd2bd8573787906",
    bold_prefix="MusicGen (genre transformation)"
)
para(
    "Both models are invoked via Replicate's REST API using an API token, and results are "
    "retrieved by polling the prediction status endpoint. Replicate charges per prediction "
    "based on GPU compute time consumed."
)

h3("5.2.4  Database — SQLite")
para(
    "SQLite is used as the application database through the SQLModel ORM. The database file "
    "(vibeshift.db) resides on the backend server. SQLite provides sufficient performance and "
    "zero-configuration setup for the current development and early-production phase. Five "
    "tables store user data, job records, stem metadata, mix outputs, and transformation jobs."
)

h3("5.2.5  Cloud Storage — AWS S3")
para(
    "All audio assets—uploaded songs, separated stems, edited stems, final mixes, and "
    "transformation outputs—are persisted in an AWS S3 bucket (vibeshift-audio-storage) in "
    "the ap-southeast-1 region. Clients receive pre-signed download URLs with a one-hour expiry, "
    "ensuring time-limited, authenticated access without exposing permanent S3 credentials."
)

h3("5.2.6  Replicate API Integration")
para(
    "VibeShift uses Replicate's hosted model infrastructure as its AI compute layer, replacing "
    "the need for managed cloud deployment, Docker orchestration, or Kubernetes clusters. "
    "Replicate provides on-demand GPU instances for each prediction request. The backend submits "
    "predictions via Replicate's REST API and polls for results, keeping the application server "
    "lightweight and stateless with respect to GPU resources."
)

# ── 5.3 ──────────────────────────────────────────────────────────────────────
h2("5.3  Frontend Implementation")

h3("5.3.1  Component Architecture")
para(
    "The frontend is structured around Expo Router's file-based routing convention, "
    "with a tab-based navigation shell and modal overlays for settings screens. The primary "
    "screens are:"
)
bullet("demixing.tsx — stem separation upload and interactive editing interface")
bullet("genre-transform.tsx — genre transformation with bubble selector and parameter controls")
bullet("library.tsx — chronological history of the user's processed audio")
bullet("profile.tsx — user profile display and statistics")
bullet("account-settings.tsx, audio-preferences.tsx, appearance.tsx — configuration screens")
para("Reusable UI components include:")
bullet("StemMixer.tsx — orchestrates the full interactive stem control interface")
bullet("StemRing.tsx — individual animated SVG ring representing one stem")
bullet("GenreBubble.tsx — floating bubble UI element for genre selection")
bullet("SongCard.tsx — card component for library song entries")
bullet("UploadButton.tsx — audio file selection and upload trigger")

h3("5.3.2  Stem Ring Visualisation")
para(
    "VibeShift replaces the conventional waveform display with a custom circular ring "
    "visualisation purpose-built for multi-stem audio editing. Each audio stem (vocals, "
    "drums, bass, other) is represented as a distinct coloured concentric ring rendered "
    "using react-native-svg. The rings surround a central disc icon and visually communicate "
    "each stem's identity through colour coding."
)
para(
    "Users interact with each ring via PanResponder gesture handlers, dragging to adjust "
    "volume, pitch shift (in semitones), and timbre strength in real time. The rings animate "
    "rotation in response to gesture input using react-native-reanimated, providing "
    "immediate tactile feedback. This design was chosen to make multi-track audio editing "
    "intuitive on a mobile touchscreen without requiring a desktop-style DAW interface."
)

h3("5.3.3  State Management")
para(
    "Application-wide state is managed through React Context. Two context providers are "
    "registered at the root level:"
)
bullet("UserContext — authentication state, user profile data, and active session token")
bullet("AppearanceContext — theming preferences (dark/light mode, accent colours)")
para(
    "The Clerk Expo SDK manages the authentication lifecycle. JWTs issued by Clerk are "
    "persisted in expo-secure-store and retrieved on each API call, attached as a Bearer "
    "token in the Authorization header. Component-local state (e.g., selected genre, "
    "stem parameter sliders, upload progress) is managed with React useState hooks."
)

h3("5.3.4  Audio Playback & Controls")
para(
    "Audio playback is implemented using expo-av. The playback interface is integrated "
    "directly with the stem editing UI, allowing users to audition the current mix while "
    "adjusting stem parameters. After modifying stem settings (pitch, timbre, volume, mute), "
    "the user triggers a re-mix request to the backend; once the new mix is returned, "
    "the player seamlessly switches to the updated audio. Standard playback controls "
    "(play, pause, seek) are rendered alongside the stem ring visualisation."
)

# ── 5.4 ──────────────────────────────────────────────────────────────────────
h2("5.4  Backend Implementation")

h3("5.4.1  API Design & Routing")
para(
    "The FastAPI application is organised into modular routers, each responsible for a "
    "distinct domain:"
)
bullet("/auth/sync (POST) — creates or updates a user record from a validated Clerk JWT")
bullet("/users/me (GET, PUT) — retrieves and updates the authenticated user's profile")
bullet(
    "/demixer/jobs (POST), /demixer/jobs/{id} (GET), "
    "/demixer/stems/{id} (PUT), /demixer/jobs/{id}/mix (POST) "
    "— full stem separation and mixing workflow"
)
bullet(
    "/transform/genres (GET), /transform/jobs (POST), /transform/jobs/{id} (GET) "
    "— genre transformation workflow"
)
bullet("/library (GET) — aggregated, date-sorted history of all user jobs")
para(
    "All endpoints return JSON-serialised Pydantic response models. HTTP status codes "
    "follow REST conventions (201 for resource creation, 200 for retrieval, 4xx for "
    "client errors, 5xx for server faults)."
)

h3("5.4.2  Authentication & Authorization (JWT)")
para(
    "Authentication is handled by Clerk. The frontend obtains a signed JWT from Clerk's "
    "Expo SDK after the user logs in. On each protected API call, this token is sent "
    "in the Authorization: Bearer header."
)
para(
    "The backend validates the token using PyJWT with the RS256 algorithm. The public "
    "key is fetched from Clerk's JWKS endpoint:"
)
code_block("https://hip-lobster-45.clerk.accounts.dev/.well-known/jwks.json")
para(
    "On successful verification, the sub claim is extracted as the clerk_user_id. "
    "FastAPI's dependency injection system (Depends()) enforces this check on every "
    "protected route. A second dependency, get_current_user(), resolves the clerk_user_id "
    "to the corresponding User row in SQLite, scoping all subsequent queries to that user."
)

h3("5.4.3  File Upload & Validation")
para(
    "Audio files are accepted as multipart/form-data requests, parsed by python-multipart. "
    "Upon receipt, the file is read into memory, content-typed, and uploaded to S3 under "
    "the path:"
)
code_block("uploads/{user_id}/{job_id}/{original_filename}")
para(
    "A corresponding job record (demix_jobs or transform_jobs) is inserted into the database "
    "with status='pending' and the S3 key stored for downstream processing. The API "
    "immediately returns the job ID to the client so polling can begin."
)

h3("5.4.4  Asynchronous Processing Pipeline")
para(
    "Job execution is handled asynchronously using FastAPI's built-in BackgroundTasks "
    "mechanism. When a job is submitted, the HTTP response is returned immediately to the "
    "client while a background coroutine begins executing the AI pipeline. The coroutine "
    "updates the job's status field in the database as processing progresses:"
)
bullet("pending → processing: AI prediction submitted to Replicate")
bullet("processing → completed: output stored in S3, download URL available to client")
bullet("processing → failed: error message recorded for client display")
para(
    "The frontend polls the job status endpoint (GET /demixer/jobs/{id} or "
    "GET /transform/jobs/{id}) at regular intervals until the status resolves. "
    "The httpx async client is used for all outbound Replicate API calls, with a "
    "600-second timeout for transformation jobs and a 120-second timeout for mix "
    "operations. This design avoids the operational complexity of a dedicated message "
    "queue while supporting fully non-blocking job execution."
)

# ── 5.5 ──────────────────────────────────────────────────────────────────────
h2("5.5  AI Microservices Implementation")

h3("5.5.1  Stem Separation — Demucs via Replicate")
para(
    "Stem separation is performed by the Demucs model hosted on Replicate "
    "(cjwbw/demucs:25a173108cff36ef9f80f854c162d01df9e6528be175794b81158fa03836d953). "
    "Demucs is a waveform-domain source separation model developed by Meta that separates "
    "a mixed audio track into up to four stems: vocals, drums, bass, and other (all "
    "remaining instruments)."
)
para("When a user submits a song for demixing, the backend executes the following steps:")
bullet("The uploaded audio file is stored in S3 and a pre-signed URL is generated.")
bullet("A prediction is submitted to Replicate's POST /v1/predictions endpoint with the S3 URL as the audio input.")
bullet("The backend polls GET /v1/predictions/{id} until the prediction status equals 'succeeded'.")
bullet("Replicate returns four separate WAV files (one per stem).")
bullet("Each stem is downloaded and re-uploaded to S3; a stems table record is created for each.")
bullet("The demix_job status is updated to 'completed' and stem download URLs are made available to the client.")

h3("5.5.2  Genre Transformation Service")
para(
    "Genre transformation uses MusicGen (meta/musicgen), a melody-conditioned music "
    "generation model from Meta, hosted on Replicate "
    "(meta/musicgen:7a76a8258b23fae65c5a22debb8841d1d7e816b75c2f24218cd2bd8573787906). "
    "The pipeline has two variants depending on whether the input song contains vocals."
)

p = doc.add_paragraph()
r = p.add_run("Instrumental-only songs:")
r.bold = True
r.underline = True
para(
    "The audio and genre text prompt are submitted directly to MusicGen on Replicate. "
    "MusicGen conditions its generation on the melody of the input audio and the target "
    "genre description, returning a transformed audio clip of the requested duration. "
    "No stem separation is required."
)

p = doc.add_paragraph()
r = p.add_run("Songs with vocals (full pipeline):")
r.bold = True
r.underline = True
para("The pipeline executes five sequential stages:")

bullet(
    "Demucs (Replicate) separates the song into a vocals track and an instrumental track "
    "(using the two-stem mode).",
    bold_prefix="Stage 1 — Stem Separation"
)
bullet(
    "MusicGen (Replicate) is applied to the instrumental track with the target genre prompt "
    "and user-specified parameters: duration (1–60 s), start_offset (seconds into the "
    "song to condition on), and guidance scale (3–15; higher values produce a stronger "
    "genre shift).",
    bold_prefix="Stage 2 — Genre Transformation of Instrumental"
)
bullet(
    "The vocals track is processed through the VocalFX module, applying genre-specific "
    "signal processing using scipy Butterworth filters. Effects per genre: "
    "Rock/Metal — tanh distortion + high-pass filter (>2 kHz); "
    "Pop — soft clipping + high-pass (>8 kHz); "
    "Jazz — low-pass filter (<8 kHz); "
    "Disco/Hip-Hop — delay echo (30 ms, 25% wet); "
    "Country/Blues — vibrato modulation (5 Hz, 0.3% depth); "
    "Reggae — delay (50 ms, 30% wet) + low-pass (<7 kHz); "
    "Classical — high-pass (>120 Hz) + subtle delay (40 ms, 15% wet).",
    bold_prefix="Stage 3 — Vocal FX"
)
bullet(
    "Beat tracking is performed on both the processed vocals and the MusicGen output "
    "using librosa.beat.beat_track(). The vocal track's tempo is adjusted to match the "
    "transformed instrumental via librosa.effects.time_stretch(), with the stretch ratio "
    "clamped to [0.7, 1.5] to prevent artefacts. The vocals are resampled to match the "
    "MusicGen output's sample rate.",
    bold_prefix="Stage 4 — Tempo Matching"
)
bullet(
    "The transformed instrumental and tempo-matched vocals are combined: "
    "output = (instrumental × instr_mix) + (vocals × vocal_mix). The result is "
    "peak-normalised to 0.95 headroom to prevent clipping and saved as the final output WAV.",
    bold_prefix="Stage 5 — Final Mix"
)

p = doc.add_paragraph()
r = p.add_run("Genre Prompt System:")
r.bold = True
para(
    "Ten genre presets are built in (blues, classical, country, disco, hiphop, jazz, metal, "
    "pop, reggae, rock), each mapped to a detailed descriptive text prompt that guides "
    "MusicGen's generation. Users may additionally supply a free-form custom prompt to target "
    "styles outside the preset list."
)

h3("5.5.3  Stem Editing — editor.py and mixer.py")
para(
    "After Demucs separation, users adjust individual stems through the frontend before "
    "requesting a final mix. Two local processing modules handle this on the backend:"
)
bullet(
    "Loads stem audio from S3. Applies pitch shifting in semitone increments using "
    "librosa.effects.pitch_shift(). Applies timbre modification via Harmonic-Percussive "
    "Source Separation (HPSS): the harmonic component is emphasised by a configurable "
    "strength factor (1.0 = no change; >1.0 = more harmonic character), enabling tonal "
    "texture adjustments without altering pitch.",
    bold_prefix="editor.py"
)
bullet(
    "Combines all stems with their current volume, pitch, and timbre settings applied. "
    "Each stem array is converted to mono, zero-padded to the length of the longest stem, "
    "summed, and normalised. The final mix is exported as a WAV file, stored in S3, and a "
    "demix_mixes record is created.",
    bold_prefix="mixer.py"
)

# ── 5.6 ──────────────────────────────────────────────────────────────────────
h2("5.6  Database Implementation")

h3("5.6.1  Schema Implementation")
para(
    "The SQLite database comprises five tables, defined via SQLModel and automatically "
    "migrated on application startup."
)
bullet(
    "Stores user profiles synchronised from Clerk. Fields: id (UUID PK), clerk_user_id "
    "(unique), email (unique), name, bio, avatar_url, audio_quality ('high' | 'standard' | "
    "'low'), export_format ('wav' | 'mp3' | 'flac'), created_at, updated_at.",
    bold_prefix="users"
)
bullet(
    "Tracks stem-separation jobs. Fields: id (UUID PK), user_id (FK → users), "
    "original_file_name, original_s3_key, song_name, duration_seconds, status "
    "('pending' | 'processing' | 'completed' | 'failed'), error_message, timestamps.",
    bold_prefix="demix_jobs"
)
bullet(
    "One row per separated stem per job. Fields: id (UUID PK), job_id (FK → demix_jobs), "
    "stem_type ('vocals' | 'drums' | 'bass' | 'other'), original_s3_key, modified_s3_key, "
    "pitch_shift (float, default 0.0), timbre_strength (float, default 1.0), volume "
    "(float, default 1.0), is_muted (boolean), timestamps.",
    bold_prefix="stems"
)
bullet(
    "Records final mixed outputs. Fields: id (UUID PK), job_id (FK → demix_jobs), "
    "output_s3_key, created_at.",
    bold_prefix="demix_mixes"
)
bullet(
    "Tracks genre transformation jobs. Fields: id (UUID PK), user_id (FK → users), "
    "original_file_name, original_s3_key, target_genre, prompt_used, output_s3_key, "
    "duration (default 10.0 s), start_offset (default 5.0 s), guidance (default 9.5), "
    "vocal_mix (default 1.5), instr_mix (default 1.0), status, error_message, timestamps.",
    bold_prefix="transform_jobs"
)

h3("5.6.2  Indexing Strategy")
para(
    "Indexes are placed on columns that appear frequently in WHERE clauses and JOIN "
    "conditions, reducing full-table scans on high-traffic query paths:"
)
bullet("users: clerk_user_id (unique index), email (unique index) — resolved on every authenticated request")
bullet("demix_jobs: user_id — used by library and status-polling queries; status — used for queue monitoring")
bullet("stems: job_id — used when loading all stems for a given job")
bullet("demix_mixes: job_id — used when retrieving the latest mix output for a job")
bullet("transform_jobs: user_id — used by library queries; target_genre — used for analytics; status — used for queue monitoring")

# ── 5.7 ──────────────────────────────────────────────────────────────────────
h2("5.7  Replicate API Integration")
para(
    "Replicate serves as the AI compute backbone of VibeShift, providing on-demand access "
    "to GPU-accelerated machine learning models through a unified REST API. This section "
    "describes how the backend integrates with Replicate for both stem separation and "
    "genre transformation."
)

h3("5.7.1  Replicate Platform Overview")
para(
    "Replicate hosts versioned, containerised ML models and exposes them via a Predictions API. "
    "Each model version is identified by a deterministic hash "
    "(model_owner/model_name:version_hash), ensuring reproducibility across calls. Compute is "
    "billed per prediction based on actual GPU wall-clock time, making it cost-efficient for "
    "workloads with variable and unpredictable request frequency. VibeShift uses two hosted "
    "models: Demucs for audio source separation and MusicGen for melody-conditioned music "
    "generation."
)

h3("5.7.2  Prediction Lifecycle")
para("Every Replicate-backed job in VibeShift follows a five-step lifecycle:")
bullet("The backend constructs a prediction request payload with the model version hash and input parameters (audio S3 URL, genre prompt, duration, guidance scale, etc.).")
bullet("A POST to Replicate's /v1/predictions endpoint creates the prediction and returns a prediction ID immediately.")
bullet("The backend enters a polling loop, calling GET /v1/predictions/{id} at intervals until the status field equals 'succeeded' or 'failed'.")
bullet("On success, the output field in the response contains one or more URLs pointing to Replicate-hosted output WAV files.")
bullet("The backend downloads each output file, stores it in S3, and updates the corresponding job record in the database to 'completed'.")

h3("5.7.3  Demucs Configuration")
para(
    "The Demucs model is called with the audio S3 URL as its primary input. Replicate returns "
    "four output URLs corresponding to the vocals, drums, bass, and other stems. For the genre "
    "transformation pipeline, the two-stem variant is used (vocals + instrumental) by passing "
    "the appropriate model parameter. The model version pinned in production is:"
)
code_block("cjwbw/demucs:25a173108cff36ef9f80f854c162d01df9e6528be175794b81158fa03836d953")

h3("5.7.4  MusicGen Configuration")
para(
    "MusicGen is called with the instrumental audio URL and the target genre text prompt. "
    "The key configurable parameters passed per request are:"
)
bullet("prompt — descriptive genre text (e.g., 'jazz music, piano, upright bass, trumpet, swing, improvisation')")
bullet("duration — output length in seconds (range: 1–60)")
bullet("guidance — classifier-free guidance scale (range: 3–15); higher values produce output that adheres more strongly to the text prompt")
bullet("model_version — set to 'melody' to enable melody-conditioning from the input audio")
para("The model version pinned in production is:")
code_block("meta/musicgen:7a76a8258b23fae65c5a22debb8841d1d7e816b75c2f24218cd2bd8573787906")

h3("5.7.5  Error Handling & Timeouts")
para(
    "If Replicate returns a prediction with status='failed', the error detail from the "
    "response is stored in the job's error_message field and the job is marked 'failed'. "
    "The frontend surfaces this message to the user. The httpx async client enforces a "
    "600-second timeout on transformation predictions and a 120-second timeout on mix "
    "operations, after which the job is marked failed with a timeout error. Replicate API "
    "credentials are stored exclusively as server-side environment variables and are never "
    "logged or transmitted to the client."
)

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = r"C:\Users\Insiya Fakhruddin\Desktop\Chapter5_Implementation.docx"
doc.save(output_path)
print(f"Saved: {output_path}")
